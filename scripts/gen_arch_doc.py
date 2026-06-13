# Generates a full architecture & flow document for the UFT CRM as a .docx.
# Run: python scripts/gen_arch_doc.py
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\Gautham\Downloads\UFT_CRM_Architecture_and_Flow.docx"

doc = Document()

# ── base styles ──
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

ACCENT = RGBColor(0x2F, 0x54, 0x96)
MUTED = RGBColor(0x60, 0x66, 0x70)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p

def p(text="", bold=False, italic=False, color=None, size=None):
    par = doc.add_paragraph()
    r = par.add_run(text)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return par

def bullet(text, level=0):
    par = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    return par

def num(text):
    return doc.add_paragraph(text, style="List Number")

def code(text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.15)
    for i, line in enumerate(text.split("\n")):
        run = par.add_run(("" if i == 0 else "\n") + line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return par

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = True
    hdr = t.rows[0].cells
    for i, hd in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(hd)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        # shade header
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "2F5496")
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
    return t

def spacer():
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# Title
# ════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("UFT CRM")
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = ACCENT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Full Architecture, Layers, APIs, Schema & End-to-End Flows")
rs.font.size = Pt(14); rs.font.color.rgb = MUTED
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run("Unitforce Technologies — Internal Sales/Manpower CRM\nTechnical Reference Document")
rm.font.size = Pt(10.5); rm.font.color.rgb = MUTED
spacer()

# ════════════════════════════════════════════════════════════════════
h("1. Executive Overview", 1)
p("UFT CRM is an internal B2B sales and manpower-recruitment CRM for Unitforce Technologies (UFT). "
  "It captures leads from multiple sources, qualifies them through an approval workflow, enriches them "
  "with third-party data, schedules meetings, and converts won opportunities into Accounts, Contacts and "
  "Deals. It layers in AI summaries (Google Gemini), role-based access control, a Naukri verification "
  "bridge for recruiters, and a fully customizable dashboard.")
p("The application is built as a single Next.js app but is internally structured into four clean layers "
  "separated by two real HTTP API seams, so the core logic and the data service can later be extracted "
  "into independent, commercially deployable services without rewriting application code.")

h("Key capabilities", 2)
for b in [
    "Lead capture, profiling (RBAC-gated), duplicate detection, and a Missing-Information queue with one-click enrichment.",
    "Response taxonomy (Cold / Neutral / Warm / Move Forward) with follow-ups, Suspend, and a Watch List.",
    "Approval workflow: Executives raise requests; Account Managers and above approve and convert to Accounts/Contacts.",
    "Meeting scheduling wizard (internal/client, online/offline) with email invites and Accept/Decline → calendar links.",
    "Multi-provider enrichment (Apollo, Hunter, PDL) and job-posting search (SerpAPI) behind a Quick Tab.",
    "AI summaries for Leads, Contacts and Accounts via Google Gemini (gemini-flash-lite-latest).",
    "Global search across all modules, CSV import/export, editable dashboard, and a shared document store.",
    "Two PostgreSQL databases (working + persistent) with reset/sync controls.",
]:
    bullet(b)

# ════════════════════════════════════════════════════════════════════
h("2. Technology Stack", 1)
table(
    ["Area", "Technology"],
    [
        ["Framework", "Next.js 16.2.6 (App Router) + React 19, TypeScript"],
        ["Styling", "Tailwind CSS 4, CSS variables for theming (Dark / Abyss / Day)"],
        ["UI libraries", "lucide-react (icons), recharts (charts), @dnd-kit (drag & drop)"],
        ["Database", "PostgreSQL 18 (node-postgres / pg driver)"],
        ["AI", "Google Gemini via Generative Language REST API"],
        ["Email", "Resend REST API"],
        ["Enrichment", "Apollo.io, Hunter.io, People Data Labs (PDL)"],
        ["Jobs data", "SerpAPI (Google Jobs engine)"],
        ["Auth/RBAC", "App-level role model (lib/permissions.ts) + per-role grants in DB"],
        ["Runtime model", "Server-authoritative; all data via /api/* route handlers"],
    ],
)

# ════════════════════════════════════════════════════════════════════
h("3. High-Level Architecture", 1)
p("Four layers separated by two HTTP seams. Dependencies flow one direction only "
  "(UI → Application API → Core → Data API → PostgreSQL). The core depends only on a Repository "
  "interface and pure contracts, never on the database driver or the framework — enforced by ESLint "
  "import boundaries.")
code(
"""┌──────────────────────────────────────────────────────────────┐
│  UI LAYER          React pages + components, contexts, hooks   │
│                    app/**/page.tsx, components/**, contexts/**  │
└───────────────┬────────────────────────────────────────────────┘
                │  Seam A — HTTP  (UI-facing, public-capable)
                ▼  fetch → /api/v1/*        (lib/api.ts, *-client.ts)
┌──────────────────────────────────────────────────────────────┐
│  APPLICATION API   app/api/v1/**  — thin controllers           │
│                    validate → call core → respond              │
└───────────────┬────────────────────────────────────────────────┘
                │  in-process call, injecting a Repository
                ▼
┌──────────────────────────────────────────────────────────────┐
│  CORE LAYER        lib/core/**  — domain services + AI modules │
│                    transport-agnostic; NO pg / NO Next imports │
└───────────────┬────────────────────────────────────────────────┘
                │  Seam B — HTTP  (internal, INTERNAL_API_KEY)
                ▼  Repository → HttpDataClient → /api/internal/v1/data/*
┌──────────────────────────────────────────────────────────────┐
│  DATA API          app/api/internal/v1/data/**                 │
│                    backed by PgRepository (the only SQL)       │
└───────────────┬────────────────────────────────────────────────┘
                ▼
         PostgreSQL  (working: uft_crm, persistent: uft_crm_persistent)"""
)
p("Transport switch: the core obtains a Repository via getRepository() (lib/data/index.ts). By default "
  "this returns an HttpDataClient that calls Seam B over HTTP (a genuine network boundary). Setting "
  "DATA_TRANSPORT=inprocess returns the in-process PgRepository instead (faster for a single process).",
  italic=True)

# ════════════════════════════════════════════════════════════════════
h("4. Layer-by-Layer Detail", 1)

h("4.1 UI Layer", 2)
p("Client-rendered React. Pages compose feature logic; shared state lives in React contexts; data access "
  "is via the useCollection hook and a set of thin client modules. The UI never imports the data layer or "
  "the pg driver (enforced).")
p("Contexts (contexts/):", bold=True)
table(["Context", "Responsibility"], [
    ["ThemeContext", "Theme selection (dark1 / dark2 / light); persisted."],
    ["NowContext", "App 'current time' (Dev Tools can override it to test time-based UI)."],
    ["CurrentUserContext", "The acting user (role drives RBAC); switchable in Dev Tools."],
    ["PermissionsContext", "canRead(module) / canWrite(module) from role grants + hierarchy."],
    ["AppDataContext", "Shared activities, follow-ups, calendar events; cascade delete + cross-entity creation."],
    ["QuickActions", "Global modals: schedule meeting, assign task, add note."],
])
p("Hooks (hooks/):", bold=True)
bullet("useCollection<T>(name) — universal CRUD over a collection (list/get/create/update/remove/replace) via lib/api.ts.")
bullet("useChartColors() — theme-aware colors for recharts.")
p("Client data modules (lib/):", bold=True)
table(["Module", "Talks to"], [
    ["lib/api.ts", "Seam A generic CRUD: /api/v1/{collection}(/{id}) + admin/db-sync."],
    ["lib/email.ts", "POST /api/v1/email/send."],
    ["lib/enrichment-client.ts", "POST /api/v1/enrich."],
    ["lib/quick-search-client.ts", "POST /api/v1/quick-search."],
    ["lib/scout-client.ts", "POST /api/v1/scout/request and /api/naukri-callback."],
    ["lib/csv.ts", "Client CSV build + download (File System Access API 'Save As')."],
])

h("4.2 Seam A — Application API (/api/v1/*)", 2)
p("Thin controllers under app/api/v1/. They validate input, call a core service (injecting "
  "getRepository()), and return JSON. This is the UI-facing, public-capable API.")
table(["Method & Path", "Purpose"], [
    ["GET /api/v1/{collection}", "List a collection."],
    ["POST /api/v1/{collection}", "Create a row."],
    ["PUT /api/v1/{collection}", "Replace an entire collection."],
    ["GET /api/v1/{collection}/{id}", "Fetch one row."],
    ["PATCH /api/v1/{collection}/{id}", "Partial update (JSONB shallow merge)."],
    ["DELETE /api/v1/{collection}/{id}", "Delete a row."],
    ["POST /api/v1/enrich", "Company/contact enrichment (Apollo/Hunter/PDL merge)."],
    ["POST /api/v1/quick-search", "Quick Tab: enrichment + job postings in parallel."],
    ["POST /api/v1/email/send", "Send email (Resend); stores meeting invite + Accept/Decline links."],
    ["POST /api/v1/scout/request", "Optional outbound webhook to the uftech.in TA module."],
    ["GET/POST /api/v1/admin/db-sync", "Persistent-DB availability + reset/save (working ↔ persistent)."],
    ["POST /api/v1/ai/lead-summary", "Gemini summary for a lead."],
    ["POST /api/v1/ai/entity-summary", "Gemini summary for a contact or account."],
])
p("Public, unversioned endpoints (stable links / external webhooks):", bold=True)
table(["Method & Path", "Purpose"], [
    ["GET /api/meeting-response", "Email Accept/Decline/ICS landing page (returns HTML/calendar)."],
    ["POST /api/naukri-callback", "Records a scout's Naukri verdict (UI + external TA module)."],
])

h("4.3 Core Layer (lib/core/**)", 2)
p("Framework- and database-agnostic domain logic. Each service receives a Repository; none import pg or "
  "next/*. Maps 1:1 to a future @uft/core package.")
table(["Module", "Responsibility"], [
    ["core/enrichment", "Orchestrates Apollo/Hunter/PDL; merges company + de-duplicates/ranks POCs (incl. phone)."],
    ["core/jobs", "Job-posting search (SerpAPI) with de-duplication."],
    ["core/email/templates", "Pure email body builders + branded HTML + ICS-ready content."],
    ["core/email/send", "Resend send + meeting-invite creation (via Repository) + Accept/Decline links."],
    ["core/calendar", "Google/Outlook calendar URLs and .ics generation (no OAuth)."],
    ["core/scout", "Naukri verdict: update scoutRequest, mirror onto lead profile, log activity."],
    ["core/meetings", "Accept/Decline invite: mark invite, log activity, mirror confirmed meeting to calendar."],
    ["core/ai", "Summarizer interface + Gemini implementation + stub fallback (Leads/Contacts/Accounts)."],
])

h("4.4 Seam B — Data API (/api/internal/v1/data/*)", 2)
p("Internal HTTP API over PostgreSQL, guarded by the INTERNAL_API_KEY header (x-internal-key). Backed "
  "directly by PgRepository — this IS the data service. Designed to be split into a standalone deployable.")
table(["Method & Path", "Purpose"], [
    ["GET/POST/PUT /api/internal/v1/data/{collection}", "List / create / replace."],
    ["GET/PATCH/DELETE /api/internal/v1/data/{collection}/{id}", "Single-row CRUD."],
    ["GET/PUT /api/internal/v1/data/db", "Whole-database snapshot / load."],
    ["GET/POST /api/internal/v1/data/admin", "Persistent configured? / reset | save."],
])

h("4.5 Data Layer (lib/data/**) & PostgreSQL", 2)
table(["File", "Responsibility"], [
    ["lib/contracts/collections.ts", "Pure types: COLLECTIONS list, CollectionName, Row, DB, emptyDB()."],
    ["lib/data/repository.ts", "The Repository interface (Seam B contract)."],
    ["lib/data/pg.ts", "PgRepository — the only place that issues SQL."],
    ["lib/data/pool.ts", "Working + persistent pg pools; lazy schema init (CREATE TABLE IF NOT EXISTS)."],
    ["lib/data/http-client.ts", "HttpDataClient — Repository over fetch to Seam B."],
    ["lib/data/index.ts", "getRepository() injector (http default, inprocess option)."],
])
p("Storage model — one table per collection:", bold=True)
code('CREATE TABLE "<collection>" (\n  id   TEXT PRIMARY KEY,\n  data JSONB NOT NULL,   -- the full row object (document-shaped)\n  seq  BIGSERIAL         -- preserves insertion order for list reads\n);')
p("Two databases: uft_crm (WORKING — live data) and uft_crm_persistent (PERSISTENT — a saved baseline). "
  "Reset restores persistent → working; Save pushes working → persistent. The persistent DB was seeded "
  "via CREATE DATABASE uft_crm_persistent TEMPLATE uft_crm.")

# ════════════════════════════════════════════════════════════════════
h("5. Data Model / Schema", 1)
p("16 collections, each stored as a JSONB document table. Field shapes below reflect the application "
  "types (optional fields shown where relevant).")

defs = [
    ("leads", "id, first_name, last_name, email, phone, company_name, source (inbound_web | n8n_apify | manual_ocr), status (new | reviewing | approved | rejected), created_at, flagged?, watchlisted?, date_of_birth?, address?, linkedin?, summary?, email_sent_at?, email_status?, profile?"),
    ("leads.profile (LeadProfile)", "industry?, company_size?, website?, open_roles?, poc_name?, poc_title?, poc_email?, poc_linkedin?, naukri_status? (pending_verification | found | not_found), naukri_url?, internal_notes?, enriched_at?, enrichment_from?, last_updated?"),
    ("contacts", "id, account_id, first_name, last_name, email, phone, job_title, account_name, date_of_birth?, address?, linkedin?, summary?, ai_summary?, flagged?"),
    ("accounts", "id, name, domain, industry, website, employee_count, annual_revenue, founded_year, contacts, deals, flagged?, ai_summary?"),
    ("deals", "id, name, stage_id, owner, account_name, total_amount, currency, created_at, contact"),
    ("products", "id, sku, name, description, base_price, billing_type (recurring | one_time), is_active"),
    ("pipelineStages", "id, name, sort_order, win_probability"),
    ("users", "id, first_name, last_name, email, role, is_active, last_login"),
    ("roles", "id, name, description, grants[] ({ module, read, write })"),
    ("activities", "id, user, entity_type, entity_name, activity_type (call_log | email | note | meeting), description, created_at, ref_id? (links to a calendar event for cascade delete)"),
    ("followUps", "id, source (lead | deal | calendar | task), source_id, entity_name, category, note?, follow_up_date, logged_at, done, assignee?, assigned_by?"),
    ("calendarEvents", "id, title, date, time?, type (meeting | task | call | deadline), assignee?, related_to?, done?, meeting_mode?, meeting_platform?, meeting_link?, location?, attendees?[], recording_url?, lead_id?"),
    ("leadRequests", "id, lead_id, lead_name, company_name, requested_by, requested_at, status (pending | approved | rejected)"),
    ("notes", "id, entity_type, entity_id, entity_name, body, author, created_at"),
    ("meetingInvites", "id, token, to, lead_id?, lead_name?, title, date, time?, mode, location?, link?, status (pending | accepted | declined), created_at, responded_at?"),
    ("scoutRequests", "id, lead_id, lead_name, company_name, poc_name?, poc_title?, poc_email?, poc_linkedin?, requested_by, assigned_to?, status (pending | found | not_found), naukri_url?, note?, requested_at, responded_at?, responded_by?"),
    ("documents", "id, name, type, date_added, modified, uploader, summary?"),
]
table(["Collection", "Fields"], defs)

# ════════════════════════════════════════════════════════════════════
h("6. RBAC & Permissions", 1)
p("Source of truth: lib/permissions.ts. A role grants Read and/or Write per module; Read controls "
  "visibility, Write controls mutating actions. Write implies Read. Grants are stored per role in the "
  "DB 'roles' collection; the hierarchy and special rules are code-defined.")
p("Modules (14):", bold=True)
p("Dashboard, Leads, Lead Profiles, Quick Search, Naukri Verification, Contacts, Accounts, Deals, "
  "Activities, Follow-ups, Business Card, Products, User Management, Roles & Permissions.")
p("Role hierarchy & rules:", bold=True)
table(["Role", "Rank", "Notes"], [
    ["Director", "4", "SUPERUSER — full access regardless of stored grants (can't be locked out)."],
    ["Business Manager", "3", "Approver/manager."],
    ["Account Manager", "2", "Approver; may delete activities; converts leads to Accounts/Contacts."],
    ["Executive", "1", "RESTRICTED — raises approval requests, sees only own activities, can't assign tasks."],
    ["Scout", "0", "Naukri verification queue (unknown roles rank 0)."],
])

# ════════════════════════════════════════════════════════════════════
h("7. Pages (app/)", 1)
table(["Route", "Module", "What it does"], [
    ["/", "Dashboard", "KPIs, charts, activity — editable panels (per-user), sections scroll vertically/horizontally."],
    ["/leads", "Leads", "Lead queue, profiling, responses, approvals, enrichment, CSV import/export, watch list, AI summary."],
    ["/contacts", "Contacts", "Contact list, detail, notes, activity timeline, AI summary."],
    ["/accounts", "Accounts", "Account cards, detail, linked contacts, notes, AI summary."],
    ["/deals", "Deals", "Kanban pipeline (drag between stages), filters, deal detail, line items on Closed Won."],
    ["/products", "Products", "Product catalog, add/edit, add-to-deal."],
    ["/activities", "Activities", "Activity log across entities (AM+ can delete)."],
    ["/follow-ups", "Follow-ups", "Follow-up queue with due dates, assignees."],
    ["/quick-tab", "Quick Search", "Company search → enrichment + contacts + job postings → add as lead."],
    ["/naukri-verify", "Naukri Verification", "Scout queue: verify leads on Naukri, record verdict."],
    ["/database", "(Database)", "Shared document store; per-doc AI summary placeholder."],
    ["/business-card", "Business Card", "Business-card scanner entry point (OCR/LLM)."],
    ["/settings", "User Mgmt / Roles", "Users & role-permission editor."],
])

h("Key components (components/)", 2)
table(["Component", "Role"], [
    ["Sidebar / TopBar / ClientLayout", "Navigation, header (search, calendar, notifications, Quick Add), provider tree."],
    ["GlobalSearch", "Search across modules; This-page vs Whole-CRM; navigate + filter + scroll + highlight."],
    ["DevTools", "Time override, user switch, settings, working↔persistent DB reset/sync."],
    ["MeetingWizard", "Internal/client + online/offline meeting scheduling with documents + email."],
    ["LogResponseModal (in leads)", "Response taxonomy, follow-ups, Suspend, Watch List, approval/convert."],
    ["CsvImportModal", "CSV upload → column mapping → inline edit → bulk import."],
    ["AiSummaryCard", "Reusable AI summary generate/regenerate card."],
    ["NotesSection / CalendarPanel / NotificationsPanel", "Per-entity notes, calendar popover, due alerts."],
    ["QuickActions / SearchableSelect / ColorFilter / NoAccess", "Global modals, typeahead select, completeness filter, access gate."],
])

# ════════════════════════════════════════════════════════════════════
h("8. End-to-End Flows", 1)

flows = [
    ("8.1 Read / list (page load)",
     ["A page mounts useCollection('leads').",
      "lib/api.ts GET /api/v1/leads (Seam A).",
      "Controller calls getRepository().list('leads').",
      "HttpDataClient GET /api/internal/v1/data/leads with x-internal-key (Seam B).",
      "PgRepository runs SELECT data FROM \"leads\" ORDER BY seq; rows return up the chain."]),
    ("8.2 Add Lead with duplicate detection",
     ["User submits the Add Lead form.",
      "findDuplicates() scans existing Leads + Contacts by email OR phone (a person may reappear with same phone, new email).",
      "If matches exist, a 'Possible duplicate' dialog lists where each lives + the matched field; user chooses Add anyway or Modify.",
      "On confirm, createLead() → POST /api/v1/leads → core/repository → Postgres."]),
    ("8.3 Log Response → follow-up / watch list / approval",
     ["User logs a response (Cold / Neutral / Warm / Move Forward).",
      "Neutral/Warm offer an 'Add to Watch List' checkbox (on by default).",
      "Categories that keep working (callback, postponed, not_responding, requested_details, meeting, suspend) create a follow-up; a 'new' lead moves to 'reviewing'.",
      "Move Forward: Executives can only Request Approval; Account Managers and above get Add-to-Accounts/Contacts options."]),
    ("8.4 Approval → convert to Accounts/Contacts",
     ["Executive raises a request (leadRequests) → appears in the Requests queue.",
      "An approver opens the request (full detail: data, notes, activity log, AI summary), then Approves.",
      "ApproveLeadModal lets the approver create an Account and/or Contact from the lead; the lead is marked approved."]),
    ("8.5 Missing Information → Get Details (enrichment)",
     ["Leads missing a critical field (name/email/phone/company/source) are held in the Missing Information tab.",
      "Get Details calls /api/v1/enrich → core/enrichment runs Apollo/Hunter/PDL in parallel.",
      "Company-level fields fill always; personal email/phone/LinkedIn fill only on a confident email/name match (no fabricated data).",
      "Once critical fields are complete the lead leaves the queue automatically."]),
    ("8.6 AI Summary (Gemini)",
     ["User clicks Generate on a Lead/Contact/Account AI Summary card.",
      "POST /api/v1/ai/(lead|entity)-summary fetches the record server-side and builds a fact list.",
      "core/ai (Gemini, gemini-flash-lite-latest) returns a 2–3 sentence summary; stored on the record (summary / ai_summary)."]),
    ("8.7 Meeting scheduling + Accept/Decline",
     ["MeetingWizard schedules a meeting and (client audience) sends an email via /api/v1/email/send.",
      "A meetingInvite row is stored with a token; the email gets Accept/Decline links to /api/meeting-response.",
      "The lead clicks Accept → invite marked accepted, activity logged, a confirmed meeting mirrored to the calendar, and add-to-calendar links (Google/Outlook/.ics) offered.",
      "Scheduling a meeting also advances a 'new' lead to 'reviewing'."]),
    ("8.8 Naukri scout verification",
     ["A Leads-write user requests verification → scoutRequest created + assigned to a Scout (+ optional webhook).",
      "The Scout queue (/naukri-verify) records Found/Not-Found with a Naukri URL.",
      "POST /api/naukri-callback → core/scout updates the request, mirrors naukri_status/naukri_url onto the lead profile, and logs an activity."]),
    ("8.9 Global search",
     ["Search lazy-loads all collections; scope is This-page or Whole-CRM.",
      "Results group by module with badges; clicking navigates to the page with ?focus=<id>.",
      "The target page applies a filter that reveals the record, scrolls it into view, and highlights it briefly."]),
    ("8.10 Reset / Sync (working ↔ persistent)",
     ["Dev Tools → Databases → Reset: /api/v1/admin/db-sync copies persistent → working (then reloads).",
      "Save: copies working → persistent (updates the baseline).",
      "Copy is an app-side snapshot/load because the two databases can't cross-query."]),
]
for title_, steps in flows:
    h(title_, 2)
    for s in steps:
        num(s)

# ════════════════════════════════════════════════════════════════════
h("9. Environment Variables", 1)
table(["Variable", "Purpose"], [
    ["DATABASE_URL", "Working Postgres (uft_crm)."],
    ["PERSISTENT_DATABASE_URL", "Persistent baseline (uft_crm_persistent). Optional — enables reset/sync."],
    ["DATA_API_URL", "Base origin of Seam B (default http://localhost:3000)."],
    ["INTERNAL_API_KEY", "Shared secret for Seam B (x-internal-key). Blank = open (local dev)."],
    ["DATA_TRANSPORT", "http (default, real seam) | inprocess (skip HTTP)."],
    ["GEMINI_API_KEY / GEMINI_MODEL", "Google AI Studio key; model default gemini-flash-lite-latest."],
    ["RESEND_API_KEY / RESEND_FROM_EMAIL", "Email sending (Resend)."],
    ["APP_BASE_URL", "Absolute base for email/meeting links (optional)."],
    ["APOLLO_API_KEY / HUNTER_API_KEY / PDL_API_KEY", "Enrichment providers (any subset)."],
    ["SERPAPI_KEY", "Job postings (Quick Tab)."],
    ["SCOUT_WEBHOOK_URL", "Optional outbound bridge to the uftech.in TA module."],
])

# ════════════════════════════════════════════════════════════════════
h("10. Build, Run & Service Extraction", 1)
p("Run locally:", bold=True)
code("npm install\nnpm run dev        # http://localhost:3000\nnpx tsc --noEmit   # type-check\nnpx eslint .       # lint (incl. layer-boundary rules)\nnpx next build     # production build")
p("Migrate legacy JSONBin data into Postgres (one-time):", bold=True)
code("node scripts/migrate-jsonbin-to-pg.mjs")
p("Extracting to independent services (commercial mode):", bold=True)
bullet("Data service: host lib/data + Seam B separately; point the app's DATA_API_URL at it and set a shared INTERNAL_API_KEY.")
bullet("Core/App API: host lib/core + app/api/v1 separately; the UI's base (/api/v1) becomes that host.")
bullet("Promote lib/contracts, lib/core, lib/data to workspace packages (@uft/contracts, @uft/core, @uft/data) — boundaries are already ESLint-enforced.")
p("Note: PostgreSQL is local-only in the current setup; a hosted Postgres (e.g. Neon/Supabase) would be "
  "required to deploy on serverless platforms such as Vercel.", italic=True, color=MUTED)

# ════════════════════════════════════════════════════════════════════
h("11. Directory Map", 1)
code(
"""app/
  page.tsx                     Dashboard (editable panels)
  leads|contacts|accounts|deals|products/page.tsx
  activities|follow-ups|settings|quick-tab|naukri-verify|database|business-card/page.tsx
  api/
    v1/                        Seam A (application API)
      [collection]/...         generic CRUD
      enrich, quick-search, email/send, scout/request, admin/db-sync
      ai/lead-summary, ai/entity-summary
    internal/v1/data/          Seam B (data API, INTERNAL_API_KEY)
      [collection]/..., db, admin
    meeting-response, naukri-callback     public endpoints
components/                    UI: TopBar, Sidebar, DevTools, GlobalSearch, MeetingWizard,
                               CsvImportModal, AiSummaryCard, NotesSection, ...
contexts/                      Theme, Now, CurrentUser, Permissions, AppData, QuickActions
hooks/                         useCollection, useChartColors
lib/
  contracts/                   pure shared types (collections, Row, DB)
  data/                        repository, pg, pool, http-client, index (getRepository)
  core/                        enrichment, jobs, email, calendar, scout, meetings, ai
  api.ts, *-client.ts, csv.ts  UI → Seam A clients
  permissions.ts, mock-data.ts, utils.ts
scripts/                       migrate-jsonbin-to-pg.mjs, gen_arch_doc.py
ARCHITECTURE.md                concise architecture reference"""
)

spacer()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = foot.add_run("UFT CRM — Technical Reference. Generated from the codebase.")
rf.italic = True; rf.font.size = Pt(9); rf.font.color.rgb = MUTED

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("Saved:", OUT)
